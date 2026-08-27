import docx

doc_path = r'c:\Users\Enrico\Desktop\Giada\Riassunti_Contabilita_Accounting.docx'
doc = docx.Document(doc_path)

for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"P{i+1}: {p.text}")

for t_idx, table in enumerate(doc.tables):
    print(f"\n--- TABLE {t_idx+1} ---")
    for r_idx, row in enumerate(table.rows):
        row_cells = [cell.text.strip() for cell in row.cells]
        print(f"Row {r_idx+1}: {' | '.join(row_cells)}")
